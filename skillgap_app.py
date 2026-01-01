import streamlit as st
import re
import PyPDF2
import io
import docx
from datetime import datetime
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Try to import OCR libraries
try:
    import fitz  # PyMuPDF
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Configure page
st.set_page_config(
    page_title="SkillGapAI",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .stApp { background: #f8f9fa; }
    h1 { color: #2c3e50; text-align: center; font-weight: 700; }
    h2, h3 { color: #34495e; }
    .stButton > button {
        background: #3498db; color: white; border: none;
        border-radius: 8px; padding: 0.6rem 1.5rem; font-weight: 600;
    }
    .stButton > button:hover { background: #2980b9; }
    [data-testid="stSidebar"] { 
        background-color: #ecf0f1;
    }
    [data-testid="stSidebar"] * {
        color: #2c3e50 !important;
    }
    [data-testid="stSidebar"] .stMarkdown {
        color: #2c3e50 !important;
    }
    [data-testid="stSidebar"] h1, [data-testid="stSidebar"] h2, [data-testid="stSidebar"] h3 {
        color: #2c3e50 !important;
    }
    [data-testid="stSidebar"] label {
        color: #2c3e50 !important;
        font-weight: 600 !important;
    }
    .milestone-box {
        padding: 15px; border-radius: 10px;
        border: 2px solid #3498db; background: white; margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

# Skill Database
SKILL_DATABASE = {
    'programming': ['python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin', 'go', 'rust', 'typescript', 'scala', 'r', 'matlab', 'perl'],
    'web_development': ['html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'spring', 'asp.net', 'jquery', 'bootstrap', 'tailwind', 'next.js', 'nuxt.js', 'svelte'],
    'mobile_development': ['android', 'ios', 'react native', 'flutter', 'xamarin', 'ionic'],
    'databases': ['sql', 'mysql', 'postgresql', 'mongodb', 'oracle', 'redis', 'cassandra', 'dynamodb', 'sqlite', 'mariadb', 'neo4j', 'elasticsearch'],
    'data_science': ['machine learning', 'deep learning', 'data analysis', 'statistics', 'pandas', 'numpy', 'scikit-learn', 'tensorflow', 'pytorch', 'keras', 'data visualization', 'tableau', 'power bi', 'matplotlib', 'seaborn'],
    'cloud': ['aws', 'azure', 'google cloud', 'gcp', 'docker', 'kubernetes', 'terraform', 'jenkins', 'ci/cd', 'devops', 'serverless', 'lambda'],
    'soft_skills': ['communication', 'leadership', 'teamwork', 'problem solving', 'critical thinking', 'project management', 'agile', 'scrum', 'time management', 'collaboration', 'presentation', 'negotiation'],
    'tools': ['git', 'github', 'gitlab', 'jira', 'confluence', 'slack', 'trello', 'postman', 'visual studio code', 'intellij', 'eclipse'],
    'methodologies': ['agile', 'scrum', 'kanban', 'waterfall', 'lean', 'six sigma', 'test-driven development', 'behavior-driven development'],
    'security': ['cybersecurity', 'encryption', 'penetration testing', 'oauth', 'jwt', 'ssl', 'firewall', 'vulnerability assessment']
}

def extract_text_from_pdf(pdf_file):
    """Extract text from uploaded PDF file including OCR for image-based PDFs"""
    text = ""
    
    try:
        # First, try normal text extraction with PyPDF2
        pdf_file.seek(0)
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_file.read()))
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        # If very little text was extracted, try OCR using PyMuPDF
        if len(text.strip()) < 50:  # Threshold for considering PDF as image-based
            if OCR_AVAILABLE:
                pdf_file.seek(0)
                
                try:
                    # Open PDF with PyMuPDF (fitz)
                    pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
                    ocr_text = ""
                    
                    # Extract images from each page and perform OCR
                    for page_num in range(len(pdf_document)):
                        page = pdf_document[page_num]
                        
                        # Try to get text first with fitz
                        page_text = page.get_text()
                        if page_text.strip():
                            ocr_text += page_text + "\n"
                        else:
                            # If no text, render page as image and do OCR
                            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better quality
                            img_data = pix.tobytes("png")
                            image = Image.open(io.BytesIO(img_data))
                            
                            # Perform OCR on the image
                            page_text = pytesseract.image_to_string(image, lang='eng')
                            if page_text.strip():
                                ocr_text += page_text + "\n"
                    
                    pdf_document.close()
                    
                    if ocr_text.strip():
                        text = ocr_text
                        
                except Exception as ocr_error:
                    # Silently handle OCR errors
                    pass
            else:
                # Silently handle missing OCR libraries
                pass
        
        return text
        
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def extract_text_from_docx(docx_file):
    """Extract text from uploaded Word document"""
    try:
        doc = docx.Document(io.BytesIO(docx_file.read()))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading Word document: {str(e)}")
        return ""

def extract_skills(text):
    """Extract skills from text using NLP techniques"""
    text_lower = text.lower()
    words = re.findall(r'\b[a-z][a-z+#.]+\b', text_lower)
    bigrams = [' '.join([words[i], words[i+1]]) for i in range(len(words)-1)]
    trigrams = [' '.join([words[i], words[i+1], words[i+2]]) for i in range(len(words)-2)]
    all_tokens = words + bigrams + trigrams
    found_skills = {}
    for category, skills in SKILL_DATABASE.items():
        found_skills[category] = []
        for skill in skills:
            count = sum(1 for token in all_tokens if token == skill.lower())
            if count > 0:
                found_skills[category].append({'skill': skill, 'count': count})
    found_skills = {k: v for k, v in found_skills.items() if v}
    return found_skills

def calculate_skill_gap(resume_skills, job_skills):
    """Calculate the gap between resume and job description skills"""
    resume_set = set()
    for skills in resume_skills.values():
        resume_set.update([s['skill'] for s in skills])
    job_set = set()
    for skills in job_skills.values():
        job_set.update([s['skill'] for s in skills])
    matching_skills = resume_set.intersection(job_set)
    missing_skills = job_set - resume_set
    additional_skills = resume_set - job_set
    if len(job_set) > 0:
        match_percentage = (len(matching_skills) / len(job_set)) * 100
    else:
        match_percentage = 0
    return {
        'matching': matching_skills,
        'missing': missing_skills,
        'additional': additional_skills,
        'match_percentage': match_percentage
    }

def get_learning_resources(skill):
    """Get learning resource recommendations for a skill"""
    return {
        'online_courses': f"Search for '{skill}' courses on Coursera, Udemy, or edX",
        'documentation': f"Official {skill} documentation and tutorials",
        'practice': f"LeetCode, HackerRank, or project-based learning for {skill}",
        'communities': f"Join {skill} communities on Reddit, Stack Overflow, or Discord"
    }

def extract_job_title(job_text):
    """Extract job title from job description"""
    lines = job_text.split('\n')
    for line in lines[:10]:
        line = line.strip()
        if line and len(line) < 100:
            if any(keyword in line.lower() for keyword in ['engineer', 'developer', 'manager', 'analyst', 'designer', 'architect', 'specialist', 'consultant', 'lead', 'senior', 'junior']):
                return line
    return "Job Position"

def create_skill_chart(gap_analysis):
    """Create a bar chart for skill analysis"""
    fig, ax = plt.subplots(figsize=(10, 6))
    categories = ['Matching Skills', 'Missing Skills', 'Additional Skills']
    values = [len(gap_analysis['matching']), len(gap_analysis['missing']), len(gap_analysis['additional'])]
    colors = ['#2ecc71', '#e74c3c', '#3498db']
    bars = ax.bar(categories, values, color=colors, alpha=0.7, edgecolor='black', linewidth=2)
    for bar in bars:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height, f'{int(height)}', ha='center', va='bottom', fontsize=12, fontweight='bold')
    ax.set_ylabel('Number of Skills', fontsize=12, fontweight='bold')
    ax.set_title('Skill Gap Analysis Overview', fontsize=14, fontweight='bold', pad=20)
    ax.grid(axis='y', alpha=0.3, linestyle='--')
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close()
    return buf

def set_cell_border(cell):
    """Set cell borders in Word document"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), '12')
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), '3498db')
        tcBorders.append(edge_element)
    tcPr.append(tcBorders)

def generate_word_report(job_title, resume_skills, job_skills, gap_analysis):
    """Generate Word document report with borders and proper formatting"""
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run("🎯 SkillGapAI Analysis Report")
    header_run.font.size = Pt(24)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(102, 126, 234)
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(job_title)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(118, 75, 162)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    
    heading = doc.add_paragraph()
    heading_run = heading.add_run("📊 Match Score")
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(44, 62, 80)
    
    metrics_table = doc.add_table(rows=2, cols=3)
    metrics_table.style = 'Light Grid Accent 1'
    headers = ['Overall Match', 'Matching Skills', 'Missing Skills']
    values = [f"{gap_analysis['match_percentage']:.1f}%", str(len(gap_analysis['matching'])), str(len(gap_analysis['missing']))]
    
    for idx, (header, value) in enumerate(zip(headers, values)):
        cell = metrics_table.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell)
        
        cell = metrics_table.rows[1].cells[idx]
        cell.text = value
        cell.paragraphs[0].runs[0].font.size = Pt(18)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(52, 152, 219)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_border(cell)
    
    doc.add_paragraph()
    
    heading = doc.add_paragraph()
    heading_run = heading.add_run("📄 Skills in Resume")
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(44, 62, 80)
    
    resume_skills_list = []
    for category, skills in resume_skills.items():
        for skill_data in skills:
            resume_skills_list.append(skill_data['skill'])
    
    skills_para = doc.add_paragraph()
    for skill in sorted(resume_skills_list):
        run = skills_para.add_run(f" {skill} ")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(21, 101, 192)
        skills_para.add_run(" | ")
    
    doc.add_paragraph()
    
    heading = doc.add_paragraph()
    heading_run = heading.add_run("💼 Skills Required in Job Description")
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(44, 62, 80)
    
    job_skills_list = []
    for category, skills in job_skills.items():
        for skill_data in skills:
            job_skills_list.append(skill_data['skill'])
    
    skills_para = doc.add_paragraph()
    for skill in sorted(job_skills_list):
        run = skills_para.add_run(f" {skill} ")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(230, 81, 0)
        skills_para.add_run(" | ")
    
    doc.add_paragraph()
    
    heading = doc.add_paragraph()
    heading_run = heading.add_run("✅ Matching Skills")
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(44, 62, 80)
    
    if gap_analysis['matching']:
        skills_para = doc.add_paragraph()
        for skill in sorted(gap_analysis['matching']):
            run = skills_para.add_run(f" ✓ {skill} ")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(21, 87, 36)
            run.font.bold = True
            skills_para.add_run(" | ")
    else:
        doc.add_paragraph("No matching skills found.")
    
    doc.add_paragraph()
    
    heading = doc.add_paragraph()
    heading_run = heading.add_run("❌ Missing Skills (Need to Learn)")
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(44, 62, 80)
    
    if gap_analysis['missing']:
        skills_para = doc.add_paragraph()
        for skill in sorted(gap_analysis['missing']):
            run = skills_para.add_run(f" ✗ {skill} ")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(114, 28, 36)
            run.font.bold = True
            skills_para.add_run(" | ")
    else:
        success_para = doc.add_paragraph("You have all the required skills!")
        success_para.runs[0].font.color.rgb = RGBColor(40, 167, 69)
        success_para.runs[0].font.bold = True
    
    doc.add_paragraph()
    doc.add_page_break()
    
    heading = doc.add_paragraph()
    heading_run = heading.add_run("📊 Skill Gap Visualization")
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(44, 62, 80)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    chart_buf = create_skill_chart(gap_analysis)
    doc.add_picture(chart_buf, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_page_break()
    
    heading = doc.add_paragraph()
    heading_run = heading.add_run("📚 Learning Recommendations")
    heading_run.font.size = Pt(16)
    heading_run.font.bold = True
    heading_run.font.color.rgb = RGBColor(44, 62, 80)
    
    if gap_analysis['missing']:
        priority_skills = list(gap_analysis['missing'])[:6]
        for idx, skill in enumerate(priority_skills, 1):
            skill_heading = doc.add_paragraph()
            skill_heading_run = skill_heading.add_run(f"{idx}. {skill.title()}")
            skill_heading_run.font.size = Pt(12)
            skill_heading_run.font.bold = True
            skill_heading_run.font.color.rgb = RGBColor(52, 152, 219)
            
            resources = get_learning_resources(skill)
            resource_table = doc.add_table(rows=4, cols=1)
            resource_table.style = 'Light List Accent 1'
            resource_data = [
                f"🎓 Online Courses: {resources['online_courses']}",
                f"📖 Documentation: {resources['documentation']}",
                f"💻 Practice: {resources['practice']}",
                f"👥 Communities: {resources['communities']}"
            ]
            
            for row_idx, resource_text in enumerate(resource_data):
                cell = resource_table.rows[row_idx].cells[0]
                cell.text = resource_text
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                set_cell_border(cell)
            doc.add_paragraph()
        
        doc.add_paragraph()
        learning_heading = doc.add_paragraph()
        learning_heading_run = learning_heading.add_run("📅 Suggested 8-Week Learning Path")
        learning_heading_run.font.size = Pt(14)
        learning_heading_run.font.bold = True
        learning_heading_run.font.color.rgb = RGBColor(102, 126, 234)
        
        learning_steps = [
            "Week 1-2: Focus on the top 2 priority skills",
            "Week 3-4: Learn the next 2-3 skills",
            "Week 5-6: Build projects incorporating new skills",
            "Week 7-8: Refine and add to your resume/portfolio"
        ]
        for step in learning_steps:
            para = doc.add_paragraph(step, style='List Bullet')
            para.runs[0].font.size = Pt(11)
    else:
        success_para = doc.add_paragraph("🌟 Congratulations! You have all the required skills for this position.")
        success_para.runs[0].font.color.rgb = RGBColor(40, 167, 69)
        success_para.runs[0].font.bold = True
        success_para.runs[0].font.size = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("SkillGapAI - Empowering Career Growth Through AI")
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    tip = doc.add_paragraph()
    tip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tip_run = tip.add_run("💡 Keep your resume updated with new skills as you learn them!")
    tip_run.font.size = Pt(9)
    tip_run.font.color.rgb = RGBColor(150, 150, 150)
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# SIDEBAR
st.sidebar.title("🎯 SkillGapAI")
st.sidebar.markdown("---")
page = st.sidebar.radio("Navigation", ["📱 Full Application", "🔸 Milestone 1", "🔸 Milestone 2", "🔸 Milestone 3", "🔸 Milestone 4"], index=0)
st.sidebar.markdown("---")
st.sidebar.markdown("### About")
st.sidebar.info("SkillGapAI helps you identify skill gaps between your resume and job descriptions, providing personalized recommendations for career growth.")
st.sidebar.markdown("---")
st.sidebar.markdown("### Milestones Overview")
st.sidebar.markdown("**Milestone 1:** Data Ingestion & Parsing\n\n**Milestone 2:** Skill Extraction\n\n**Milestone 3:** Skill Gap Analysis\n\n**Milestone 4:** Dashboard & Reporting")

# MAIN CONTENT
if page == "📱 Full Application":
    st.title("🎯 SkillGapAI - Full Application")
    st.markdown("### Analyze Resume and Job Description for Skill Gaps")
    st.markdown("---")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("📄 Resume")
        resume_file = st.file_uploader("Upload Resume (PDF only)", type=['pdf'], key="resume_pdf")
        resume_text = ""
        if resume_file:
            with st.spinner("📄 Extracting text from resume..."):
                resume_text = extract_text_from_pdf(resume_file)
            if resume_text:
                st.success("✅ Resume uploaded and processed successfully!")
                with st.expander("Preview extracted text"):
                    st.text_area("Resume Text", resume_text[:500] + "...", height=150, disabled=True)
            else:
                st.error("❌ Could not extract text from resume. Please try a different file.")

    with col2:
        st.subheader("💼 Job Description")
        job_input_type = st.radio("Input Type:", ["Text", "Upload File (PDF/Word/TXT)"], key="job")
        job_text = ""
        if job_input_type == "Text":
            job_text = st.text_area("Paste job description here:", height=300, key="job_text")
        else:
            job_file = st.file_uploader("Upload Job Description", type=['pdf', 'docx', 'txt'], key="job_file")
            if job_file:
                file_type = job_file.name.split('.')[-1].lower()
                with st.spinner(f"📄 Processing {file_type.upper()} file..."):
                    if file_type == 'pdf':
                        job_text = extract_text_from_pdf(job_file)
                    elif file_type == 'docx':
                        job_text = extract_text_from_docx(job_file)
                    elif file_type == 'txt':
                        job_text = job_file.read().decode('utf-8')
                if job_text:
                    st.success(f"✅ Job description uploaded successfully!")
                    with st.expander("Preview extracted text"):
                        st.text_area("Job Description Text", job_text[:500] + "...", height=150, disabled=True)
                else:
                    st.error("❌ Could not extract text. Please try a different file.")

    st.markdown("---")
    if st.button("🔍 Analyze Skill Gap", type="primary", width="stretch"):
        if resume_text and job_text:
            with st.spinner("Analyzing skills and identifying gaps..."):
                job_title = extract_job_title(job_text)
                resume_skills = extract_skills(resume_text)
                job_skills = extract_skills(job_text)
                gap_analysis = calculate_skill_gap(resume_skills, job_skills)
                
                st.session_state['analysis_complete'] = True
                st.session_state['job_title'] = job_title
                st.session_state['resume_skills'] = resume_skills
                st.session_state['job_skills'] = job_skills
                st.session_state['gap_analysis'] = gap_analysis
                
                st.markdown("---")
                st.header("📊 Analysis Results")
                
                match_pct = gap_analysis['match_percentage']
                col_metric1, col_metric2, col_metric3 = st.columns(3)
                with col_metric1:
                    st.metric("Match Score", f"{match_pct:.1f}%")
                with col_metric2:
                    st.metric("Matching Skills", len(gap_analysis['matching']))
                with col_metric3:
                    st.metric("Missing Skills", len(gap_analysis['missing']))
                
                if match_pct >= 80:
                    feedback = "🎉 Excellent match! You're well-qualified for this role."
                elif match_pct >= 60:
                    feedback = "👍 Good match! Consider learning a few more skills to strengthen your application."
                elif match_pct >= 40:
                    feedback = "⚠️ Moderate match. Upskilling in missing areas is recommended."
                else:
                    feedback = "📚 Significant gap detected. Focus on building the core required skills."
                
                st.progress(match_pct / 100)
                st.info(feedback)
                
                tab1, tab2, tab3, tab4 = st.tabs(["✅ Matching Skills", "❌ Missing Skills", "➕ Additional Skills", "📚 Recommendations"])
                
                with tab1:
                    st.subheader("Skills You Already Have")
                    if gap_analysis['matching']:
                        for category, skills in resume_skills.items():
                            matching_in_category = [s['skill'] for s in skills if s['skill'] in gap_analysis['matching']]
                            if matching_in_category:
                                st.markdown(f"**{category.replace('_', ' ').title()}:**")
                                cols = st.columns(4)
                                for idx, skill in enumerate(matching_in_category):
                                    with cols[idx % 4]:
                                        st.success(f"✓ {skill}")
                    else:
                        st.warning("No matching skills found.")
                
                with tab2:
                    st.subheader("Skills You Need to Develop")
                    if gap_analysis['missing']:
                        for category, skills in job_skills.items():
                            missing_in_category = [s['skill'] for s in skills if s['skill'] in gap_analysis['missing']]
                            if missing_in_category:
                                st.markdown(f"**{category.replace('_', ' ').title()}:**")
                                cols = st.columns(4)
                                for idx, skill in enumerate(missing_in_category):
                                    with cols[idx % 4]:
                                        st.error(f"✗ {skill}")
                    else:
                        st.success("You have all the required skills! 🎉")
                
                with tab3:
                    st.subheader("Additional Skills You Bring")
                    if gap_analysis['additional']:
                        st.info("These skills weren't explicitly mentioned in the job description but add value to your profile.")
                        cols = st.columns(4)
                        for idx, skill in enumerate(sorted(gap_analysis['additional'])):
                            with cols[idx % 4]:
                                st.info(f"• {skill}")
                    else:
                        st.warning("No additional skills beyond job requirements.")
                
                with tab4:
                    st.subheader("Upskilling Recommendations")
                    if gap_analysis['missing']:
                        st.markdown("### 🎯 Priority Skills to Learn")
                        priority_skills = list(gap_analysis['missing'])[:6]
                        for idx, skill in enumerate(priority_skills, 1):
                            with st.expander(f"{idx}. {skill.title()}", expanded=(idx <= 3)):
                                resources = get_learning_resources(skill)
                                st.markdown(f"**🎓 Online Courses:** {resources['online_courses']}")
                                st.markdown(f"**📖 Documentation:** {resources['documentation']}")
                                st.markdown(f"**💻 Practice:** {resources['practice']}")
                                st.markdown(f"**👥 Communities:** {resources['communities']}")
                        st.markdown("---")
                        st.markdown("### 📅 Suggested Learning Path")
                        st.markdown("""
                        1. **Week 1-2:** Focus on the top 2 priority skills
                        2. **Week 3-4:** Learn the next 2-3 skills
                        3. **Week 5-6:** Build projects incorporating new skills
                        4. **Week 7-8:** Refine and add to your resume/portfolio
                        """)
                    else:
                        st.success("🌟 You have all required skills! Consider deepening your expertise in your strongest areas.")
        else:
            st.error("⚠️ Please provide both resume and job description to analyze.")

    if st.session_state.get('analysis_complete', False):
        st.markdown("---")
        st.markdown("### 📥 Download Analysis Report")
        col_download1, col_download2 = st.columns(2)
        with col_download1:
            if st.button("📄 Download Word Report (.docx)", type="primary", width="stretch"):
                with st.spinner("Generating Word report..."):
                    word_report = generate_word_report(st.session_state['job_title'], st.session_state['resume_skills'], st.session_state['job_skills'], st.session_state['gap_analysis'])
                    st.download_button(label="⬇️ Download Word Report", data=word_report, file_name=f"SkillGap_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", width="stretch")
                    st.success("✅ Word report generated! Click 'Download Word Report' button above.")

elif page == "🔸 Milestone 1":
    st.title("🔸 Milestone 1: Data Ingestion & Parsing")
    st.markdown("---")
    st.markdown("""<div class='milestone-box'><h3>🎯 Objective</h3><p>Collect resume and job description data through file upload or text input, extract readable text from various formats (PDF, DOCX, TXT), and clean and preview the content for further analysis.</p></div>""", unsafe_allow_html=True)
    st.markdown("### 📤 Upload Documents")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("📄 Resume Upload")
        resume_file_m1 = st.file_uploader("Upload Resume (PDF)", type=['pdf'], key="resume_m1")
        if resume_file_m1:
            with st.spinner("📄 Extracting text from resume..."):
                resume_text_m1 = extract_text_from_pdf(resume_file_m1)
            if resume_text_m1:
                st.success("✅ Resume processed successfully!")
                st.session_state['resume_text_m1'] = resume_text_m1
                st.session_state['m1_resume_uploaded'] = True
                st.markdown("#### 👁️ Preview Extracted Text")
                st.text_area("Resume Content", resume_text_m1, height=300, key="resume_preview_m1")
                st.metric("Characters Extracted", len(resume_text_m1))
                st.metric("Words Count", len(resume_text_m1.split()))
    
    with col2:
        st.subheader("💼 Job Description Upload")
        job_input_type_m1 = st.radio("Input Type:", ["Text", "Upload File"], key="job_m1_type")
        job_text_m1 = ""
        
        if job_input_type_m1 == "Text":
            job_text_m1 = st.text_area("Paste job description:", height=300, key="job_text_m1_input")
            if job_text_m1:
                st.session_state['job_text_m1'] = job_text_m1
                st.session_state['m1_job_uploaded'] = True
                st.success("✅ Job description saved!")
        else:
            job_file_m1 = st.file_uploader("Upload Job Description", type=['pdf', 'docx', 'txt'], key="job_file_m1")
            if job_file_m1:
                file_type = job_file_m1.name.split('.')[-1].lower()
                with st.spinner(f"📄 Processing {file_type.upper()} file..."):
                    if file_type == 'pdf':
                        job_text_m1 = extract_text_from_pdf(job_file_m1)
                    elif file_type == 'docx':
                        job_text_m1 = extract_text_from_docx(job_file_m1)
                    elif file_type == 'txt':
                        job_text_m1 = job_file_m1.read().decode('utf-8')
                
                if job_text_m1:
                    st.success("✅ Job description processed successfully!")
                    st.session_state['job_text_m1'] = job_text_m1
                    st.session_state['m1_job_uploaded'] = True
                    st.markdown("#### 👁️ Preview Extracted Text")
                    st.text_area("Job Description Content", job_text_m1, height=300, key="job_preview_m1")
                    st.metric("Characters Extracted", len(job_text_m1))
                    st.metric("Words Count", len(job_text_m1.split()))
    
    st.markdown("---")
    if st.session_state.get('m1_resume_uploaded', False) and st.session_state.get('m1_job_uploaded', False):
        st.success("✅ **Milestone 1 Complete:** Both documents have been uploaded. Proceed to Milestone 2 for skill extraction.")
    else:
        st.info("📝 Please upload both resume and job description to complete Milestone 1.")

elif page == "🔸 Milestone 2":
    st.title("🔸 Milestone 2: Skill Extraction")
    st.markdown("---")
    st.markdown("""<div class='milestone-box'><h3>🎯 Objective</h3><p>Identify and extract technical and soft skills from both the resume and job description using NLP techniques and predefined skill dictionaries. Display structured skill lists for easy comparison.</p></div>""", unsafe_allow_html=True)
    
    # Check if data exists from Milestone 1
    has_data = 'resume_text_m1' in st.session_state and 'job_text_m1' in st.session_state
    
    if not has_data:
        st.warning("⚠️ No data found from Milestone 1. Please upload documents below.")
        st.markdown("### 📤 Upload Documents")
        
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("📄 Resume Upload")
            resume_file_m2 = st.file_uploader("Upload Resume (PDF)", type=['pdf'], key="resume_m2_direct")
            if resume_file_m2:
                with st.spinner("Processing resume..."):
                    resume_text = extract_text_from_pdf(resume_file_m2)
                    if resume_text:
                        st.session_state['resume_text_m1'] = resume_text
                        st.success("✅ Resume uploaded successfully!")
                        st.rerun()
        
        with col2:
            st.subheader("💼 Job Description Upload")
            job_file_m2 = st.file_uploader("Upload Job Description", type=['pdf', 'docx', 'txt'], key="job_m2_direct")
            if job_file_m2:
                file_type = job_file_m2.name.split('.')[-1].lower()
                with st.spinner("Processing job description..."):
                    if file_type == 'pdf':
                        job_text = extract_text_from_pdf(job_file_m2)
                    elif file_type == 'docx':
                        job_text = extract_text_from_docx(job_file_m2)
                    elif file_type == 'txt':
                        job_text = job_file_m2.read().decode('utf-8')
                    if job_text:
                        st.session_state['job_text_m1'] = job_text
                        st.success("✅ Job description uploaded successfully!")
                        st.rerun()
    else:
        st.success("✅ Data loaded from Milestone 1")
        
        # Show data info
        with st.expander("📄 View Resume Data"):
            st.text_area("Resume Content", st.session_state['resume_text_m1'][:500] + "...", height=150, disabled=True, key="resume_preview_m2")
        
        with st.expander("💼 View Job Description Data"):
            st.text_area("Job Description Content", st.session_state['job_text_m1'][:500] + "...", height=150, disabled=True, key="job_preview_m2")
        
        st.markdown("---")
        
        # Extract skills button
        if st.button("🔍 Extract Skills", type="primary", width="stretch", key="extract_skills_btn"):
            with st.spinner("Extracting skills using NLP..."):
                resume_skills_m2 = extract_skills(st.session_state['resume_text_m1'])
                job_skills_m2 = extract_skills(st.session_state['job_text_m1'])
                st.session_state['resume_skills_m2'] = resume_skills_m2
                st.session_state['job_skills_m2'] = job_skills_m2
                st.session_state['skills_extracted'] = True
                st.rerun()
        
        # Display results if skills have been extracted
        if st.session_state.get('skills_extracted', False):
            st.markdown("---")
            st.header("📊 Extracted Skills")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("📄 Skills Found in Resume")
                if 'resume_skills_m2' in st.session_state:
                    resume_skills_m2 = st.session_state['resume_skills_m2']
                    total_resume_skills = 0
                    for category, skills in resume_skills_m2.items():
                        st.markdown(f"**{category.replace('_', ' ').title()}** ({len(skills)} skills)")
                        for skill_data in skills:
                            st.markdown(f"- {skill_data['skill']} (mentioned {skill_data['count']}x)")
                        total_resume_skills += len(skills)
                        st.markdown("")
                    st.metric("Total Skills in Resume", total_resume_skills)
            
            with col2:
                st.subheader("💼 Skills Required in Job Description")
                if 'job_skills_m2' in st.session_state:
                    job_skills_m2 = st.session_state['job_skills_m2']
                    total_job_skills = 0
                    for category, skills in job_skills_m2.items():
                        st.markdown(f"**{category.replace('_', ' ').title()}** ({len(skills)} skills)")
                        for skill_data in skills:
                            st.markdown(f"- {skill_data['skill']} (mentioned {skill_data['count']}x)")
                        total_job_skills += len(skills)
                        st.markdown("")
                    st.metric("Total Skills in Job Description", total_job_skills)
            
            st.markdown("---")
            st.info("✅ **Milestone 2 Complete:** Skills extracted successfully. Proceed to Milestone 3 for gap analysis.")

elif page == "🔸 Milestone 3":
    st.title("🔸 Milestone 3: Skill Gap Analysis")
    st.markdown("---")
    st.markdown("""<div class='milestone-box'><h3>🎯 Objective</h3><p>Compare extracted skills to identify matched skills (skills you have that are required), missing skills (skills you need to learn), and extra skills (additional skills you bring). Generate personalized improvement recommendations.</p></div>""", unsafe_allow_html=True)
    
    has_milestone2_data = 'resume_skills_m2' in st.session_state and 'job_skills_m2' in st.session_state
    
    if has_milestone2_data:
        st.success("✅ Skills loaded from Milestone 2")
        
        # Show preview of skills count
        st.markdown("### 📊 Skills Summary")
        col_preview1, col_preview2 = st.columns(2)
        with col_preview1:
            total_resume = sum(len(skills) for skills in st.session_state['resume_skills_m2'].values())
            st.metric("Resume Skills", total_resume)
        with col_preview2:
            total_job = sum(len(skills) for skills in st.session_state['job_skills_m2'].values())
            st.metric("Job Description Skills", total_job)
        
        st.markdown("---")
        
        if st.button("📊 Analyze Skill Gap", type="primary", width="stretch", key="analyze_gap_btn"):
            with st.spinner("Analyzing skill gaps..."):
                gap_analysis_m3 = calculate_skill_gap(st.session_state['resume_skills_m2'], st.session_state['job_skills_m2'])
                st.session_state['gap_analysis_m3'] = gap_analysis_m3
                st.session_state['gap_analyzed'] = True
                st.rerun()
        
        # Display results if gap has been analyzed
        if st.session_state.get('gap_analyzed', False):
            gap_analysis_m3 = st.session_state['gap_analysis_m3']
            
            st.markdown("---")
            st.header("📈 Gap Analysis Results")
            
            match_pct = gap_analysis_m3['match_percentage']
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Match Score", f"{match_pct:.1f}%")
            with col2:
                st.metric("Matching Skills", len(gap_analysis_m3['matching']))
            with col3:
                st.metric("Missing Skills", len(gap_analysis_m3['missing']))
            
            st.progress(match_pct / 100)
            
            st.markdown("---")
            st.subheader("✅ Matching Skills")
            if gap_analysis_m3['matching']:
                cols = st.columns(4)
                for idx, skill in enumerate(sorted(gap_analysis_m3['matching'])):
                    with cols[idx % 4]:
                        st.success(f"✓ {skill}")
            else:
                st.warning("No matching skills found.")
            
            st.markdown("---")
            st.subheader("❌ Missing Skills")
            if gap_analysis_m3['missing']:
                cols = st.columns(4)
                for idx, skill in enumerate(sorted(gap_analysis_m3['missing'])):
                    with cols[idx % 4]:
                        st.error(f"✗ {skill}")
            else:
                st.success("No missing skills! You have everything required.")
            
            st.markdown("---")
            st.subheader("➕ Additional Skills")
            if gap_analysis_m3['additional']:
                st.info("Skills not mentioned in job description but present in your resume")
                cols = st.columns(4)
                for idx, skill in enumerate(sorted(gap_analysis_m3['additional'])):
                    with cols[idx % 4]:
                        st.info(f"• {skill}")
            
            st.markdown("---")
            st.subheader("📚 Recommendations")
            if gap_analysis_m3['missing']:
                st.markdown("**Priority Skills to Learn:**")
                priority_skills = list(gap_analysis_m3['missing'])[:5]
                for idx, skill in enumerate(priority_skills, 1):
                    st.markdown(f"{idx}. **{skill.title()}**")
                    resources = get_learning_resources(skill)
                    st.markdown(f"   - 🎓 {resources['online_courses']}")
                    st.markdown(f"   - 📖 {resources['documentation']}")
            else:
                st.success("🌟 You have all the required skills!")
            
            st.markdown("---")
            st.info("✅ **Milestone 3 Complete:** Gap analysis done. Proceed to Milestone 4 for visualization.")
    else:
        st.warning("⚠️ Please complete Milestone 2 first to extract skills.")
        st.info("💡 Navigate to **Milestone 1** to upload documents, then **Milestone 2** to extract skills.")

elif page == "🔸 Milestone 4":
    st.title("🔸 Milestone 4: Dashboard & Reporting")
    st.markdown("---")
    st.markdown("""<div class='milestone-box'><h3>🎯 Objective</h3><p>Visualize skill comparison using interactive charts and graphs for quick understanding of gaps and strengths. Provide a comprehensive dashboard view of the analysis results.</p></div>""", unsafe_allow_html=True)
    
    has_milestone3_data = 'gap_analysis_m3' in st.session_state
    
    if has_milestone3_data:
        st.success("✅ Analysis data loaded from Milestone 3")
        gap_analysis_m4 = st.session_state['gap_analysis_m3']
        st.markdown("---")
        st.subheader("📊 Skill Gap Visualization")
        chart_buf = create_skill_chart(gap_analysis_m4)
        st.image(chart_buf, width="stretch")
        st.markdown("---")
        st.subheader("📈 Key Metrics Dashboard")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Match Percentage", f"{gap_analysis_m4['match_percentage']:.1f}%", delta=f"{gap_analysis_m4['match_percentage'] - 50:.1f}%" if gap_analysis_m4['match_percentage'] >= 50 else None)
        with col2:
            st.metric("Matching Skills", len(gap_analysis_m4['matching']), delta="Good" if len(gap_analysis_m4['matching']) > 5 else None)
        with col3:
            st.metric("Missing Skills", len(gap_analysis_m4['missing']), delta="Needs Improvement" if len(gap_analysis_m4['missing']) > 5 else None, delta_color="inverse")
        with col4:
            st.metric("Additional Skills", len(gap_analysis_m4['additional']), delta="Bonus" if len(gap_analysis_m4['additional']) > 0 else None)
        st.markdown("---")
        st.subheader("📊 Skill Distribution")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Skills by Status**")
            status_data = {"Status": ["Matching", "Missing", "Additional"], "Count": [len(gap_analysis_m4['matching']), len(gap_analysis_m4['missing']), len(gap_analysis_m4['additional'])]}
            st.bar_chart(status_data, x="Status", y="Count")
        with col2:
            st.markdown("**Match Quality Indicator**")
            match_pct = gap_analysis_m4['match_percentage']
            if match_pct >= 80:
                st.success("🌟 Excellent Match (80%+)")
                st.markdown("You're highly qualified for this position!")
            elif match_pct >= 60:
                st.info("👍 Good Match (60-79%)")
                st.markdown("You meet most requirements. A few skills to improve.")
            elif match_pct >= 40:
                st.warning("⚠️ Moderate Match (40-59%)")
                st.markdown("Consider upskilling in key areas.")
            else:
                st.error("📚 Low Match (<40%)")
                st.markdown("Significant learning required for this role.")
        st.markdown("---")
        st.info("✅ **Milestone 4 Complete:** Dashboard and visualizations ready. Analysis complete!")
    else:
        st.warning("⚠️ Please complete Milestone 3 first to perform gap analysis.")
        st.info("💡 Navigate to Milestone 1 → 2 → 3 to complete the analysis pipeline.")

st.markdown("---")
st.markdown("""<div style='text-align: center; padding: 20px;'><p style='color: #555; font-size: 1rem;'>SkillGapAI - Empowering Career Growth Through AI</p><p style='color: #888; font-size: 0.9rem;'>💡 Tip: Keep your resume updated with new skills as you learn them!</p></div>""", unsafe_allow_html=True)